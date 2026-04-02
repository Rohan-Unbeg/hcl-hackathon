"""
Core Extraction Engine.
Handles PDF, DOCX, and Image (OCR) processing with layout preservation.
"""

import base64
import io
import logging
from typing import Dict, Any

import fitz  # PyMuPDF
from docx import Document
from PIL import Image, ImageOps
import pytesseract

logger = logging.getLogger(__name__)


def process_document(file_base64: str, file_type: str) -> str:
    """
    Decodes and extracts text content from various document formats.
    """
    try:
        file_bytes = base64.b64decode(file_base64)
    except Exception as e:
        logger.error(f"Base64 decoding failed: {e}")
        raise ValueError("Invalid base64 encoding for document.")

    extractors = {
        "pdf": _extract_pdf,
        "docx": _extract_docx,
        "image": _extract_image
    }

    if file_type.lower() not in extractors:
        raise ValueError(f"Unsupported file type: {file_type}")

    return extractors[file_type.lower()](file_bytes)


def _extract_pdf(data: bytes) -> str:
    """Extracts text from PDF while maintaining reading order."""
    text = ""
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            text += page.get_text("text") + "\n\n"
    return text.strip()


def _extract_docx(data: bytes) -> str:
    """Extracts text from paragraphs and tables in DOCX."""
    doc = Document(io.BytesIO(data))
    full_text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))

    return "\n".join(full_text).strip()


def _extract_image(data: bytes) -> str:
    """Performs OCR on images with pre-processing for better accuracy."""
    img = Image.open(io.BytesIO(data))

    # Pre-processing for OCR accuracy
    img = ImageOps.grayscale(img)
    img = ImageOps.autocontrast(img)

    # OCR with high-quality settings
    text = pytesseract.image_to_string(img, config='--oem 3 --psm 6')
    return text.strip()
